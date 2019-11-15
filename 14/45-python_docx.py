#! python3

import docx

docxObj = docx.Document(r'D:\Users\pavan\Documents\GitHub-lakshmanboddoju\Automate_the_boring_stuff_with_python\14\demo.docx')
paraObj = docObj.paragraphs[1]

print(paraObj.text)
print(paraObj.runs[0].text)
print(paraObj.runs[3].text)
print(paraObj.runs[3].bold == true)
print(paraOnj.runs[3].italic)

paraObj.runs[3].underline = True
paraObj.runs[3].text = 'italic and underlined'

docxObj.save(r'D:\Users\pavan\Documents\GitHub-lakshmanboddoju\Automate_the_boring_stuff_with_python\14\demo2.docx')

# Ctrl + Alt + Shift + S in  word to get styles
print(paraObj.style())
paraObj.style = 'Title'

docxObj.save(r'D:\Users\pavan\Documents\GitHub-lakshmanboddoju\Automate_the_boring_stuff_with_python\14\demo3.docx')

# New doc creation
newDoc = docx.Document()
newDoc.add_paragraph('Hello World!')
newDoc.add_paragraph('This is a new paragraph.')

newDoc.save(r'D:\Users\pavan\Documents\GitHub-lakshmanboddoju\Automate_the_boring_stuff_with_python\14\demo4.docx')

newParaObj = newDoc.paragraphs[0]
newParaObj.add_run('This is a new run.')
newParaObj.runs[-1].bold = True

newDoc.save(r'D:\Users\pavan\Documents\GitHub-lakshmanboddoju\Automate_the_boring_stuff_with_python\14\demo5.docx')

